import os
import shutil
from docx import Document
from docx.shared import Inches, Pt

def clean_text(text):
    return "".join(ch for ch in text if ord(ch) in [0x9, 0xA, 0xD] or 0x20 <= ord(ch) <= 0xD7FF or 0xE000 <= ord(ch) <= 0xFFFD)

def convert_md_to_docx(md_path, docx_path, blind=False):
    print(f"Reading manuscript from {md_path}...")
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    doc = Document()
    
    # Set standard margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Set default font style (Times New Roman, 12pt)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    skip_mode = False
    
    for line in lines:
        stripped = line.strip()
        
        # Blind author/affiliation details if requested
        if blind:
            if stripped.startswith("**Author:**"):
                p = doc.add_paragraph()
                p.add_run("**Author:** [Blinded for Peer Review]").bold = True
                continue
            if stripped.startswith("**Affiliation:**"):
                p = doc.add_paragraph()
                p.add_run("**Affiliation:** [Blinded for Peer Review]").bold = True
                continue
            if stripped.startswith("## Code Availability Statement"):
                skip_mode = True
                p = doc.add_paragraph()
                p.add_run("## Code Availability Statement\n[Blinded for Peer Review]").bold = True
                continue
            if skip_mode:
                if stripped.startswith("---") or stripped.startswith("## References"):
                    skip_mode = False
                else:
                    continue # Skip content under Code Availability Statement
        
        if stripped.startswith("# "):
            doc.add_heading(clean_text(stripped[2:]), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(clean_text(stripped[3:]), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(clean_text(stripped[4:]), level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(clean_text(stripped[5:]), level=4)
        elif stripped == "---":
            # Add a horizontal line/divider or page break
            doc.add_page_break()
        elif stripped == "" or stripped.isspace():
            continue
        else:
            # Clean up standard formatting (bold/italic)
            p = doc.add_paragraph()
            # Simple bold search
            run_text = stripped
            p.add_run(clean_text(run_text))
            
    doc.save(docx_path)
    print(f"Saved docx to {docx_path}")

def build_docs():
    output_dir = "submission_files"
    os.makedirs(output_dir, exist_ok=True)
    
    # 1. Generate Cover_Letter.docx
    print("Building Cover_Letter.docx...")
    cover_doc = Document()
    style = cover_doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    with open("cover_letter.md", "r", encoding="utf-8") as f:
        letter_content = f.read()
        
    # Remove markdown header
    letter_content = letter_content.replace("# Cover Letter for Manuscript Submission\n\n", "")
    for paragraph in letter_content.split("\n\n"):
        if paragraph.strip() != "":
            cover_doc.add_paragraph(paragraph.strip())
            
    cover_doc.save(os.path.join(output_dir, "Cover_Letter.docx"))
    
    # 2. Generate Title_Page.docx
    print("Building Title_Page.docx...")
    title_doc = Document()
    style = title_doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    title_doc.add_paragraph("").paragraph_format.space_before = Pt(36)
    h = title_doc.add_heading("An Explainable Machine Learning Framework for Predicting 30-Day Hospital Readmissions Using Electronic Health Records", level=1)
    h.alignment = 1 # Center
    
    title_doc.add_paragraph("").paragraph_format.space_before = Pt(24)
    p_author = title_doc.add_paragraph()
    p_author.add_run("Author: ").bold = True
    p_author.add_run("Mustaid Ahmed")
    p_author.alignment = 1
    
    p_aff = title_doc.add_paragraph()
    p_aff.add_run("Affiliation: ").bold = True
    p_aff.add_run("Saint Louis University, St. Louis, MO, USA")
    p_aff.alignment = 1
    
    p_contact = title_doc.add_paragraph()
    p_contact.add_run("Corresponding Author Email: ").bold = True
    p_contact.add_run("[Insert your email address here]")
    p_contact.alignment = 1
    
    title_doc.save(os.path.join(output_dir, "Title_Page.docx"))
    
    # 3. Generate Manuscript_Blind.docx
    convert_md_to_docx("hospital_readmission_paper.md", os.path.join(output_dir, "Manuscript_Blind.docx"), blind=True)
    
    # 4. Generate Declaration_of_Interest.txt
    print("Building Declaration_of_Interest.txt...")
    statement = (
        "Declaration of Interest Statement\n\n"
        "The author declares that they have no known competing financial interests or "
        "personal relationships that could have appeared to influence the work reported in this paper."
    )
    with open(os.path.join(output_dir, "Declaration_of_Interest.txt"), "w", encoding="utf-8") as f:
        f.write(statement)
        
    # 5. Copy Figures
    print("Copying figures to submission folder...")
    figs = [
        "roc_curves.png",
        "calibration_curves.png",
        "feature_importance.png",
        "shap_summary_plot.png",
        "shap_local_high_risk.png",
        "lime_local_high_risk.png",
        "shap_local_low_risk.png"
    ]
    for fig in figs:
        src = os.path.join("results", fig)
        dst = os.path.join(output_dir, fig)
        if os.path.exists(src):
            shutil.copy(src, dst)
            print(f"Copied {fig}")
            
    print(f"\nAll files prepared in the '{output_dir}/' folder successfully.")

if __name__ == "__main__":
    build_docs()
